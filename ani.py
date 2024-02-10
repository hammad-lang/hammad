from docx import Document
from docx.enum.text import WD_BREAK
from openai import OpenAI
from docx.shared import Pt

client = OpenAI(
    api_key="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",
)
def chat_with_openai(prompt):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}]
    )
    chatbot_response = response.choices[0].message.content
    return chatbot_response.strip()

def convert_docx_to_new_docx(original_docx_file,new_docx_file):
    doc = Document(original_docx_file)
    new_doc = Document()

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)

    for paragraph in doc.paragraphs:
        print("This is style",paragraph.style.name)

        print(paragraph.text)

        if paragraph.text=="":
            print("space")


        elif paragraph.style.name.startswith('Heading') or paragraph.style.name == "Title":
           h = new_doc.add_heading(paragraph.text)

        else:
            prompt = (f'''I am giving you text convert it to simpler english: this is the text : {paragraph.text}''')
            text=chat_with_openai(prompt)



            if paragraph.style.name.startswith('List'):
                l = new_doc.add_paragraph(text, style="List Bullet")

            elif paragraph.style.name.startswith('Normal'):
                p = new_doc.add_paragraph(text, "Normal")

        for run in paragraph.runs:
            if 'lastRenderedPageBreak' in run._element.xml:
                new_doc.add_page_break()
                print('soft page break found at run:', run.text[:20])
            if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                new_doc.add_page_break()
                print('hard page break found at run:', run.text[:20])


    new_doc.save(new_docx_file)
    remove_unwanted_spaces("output.docx")

convert_docx_to_new_docx("animation_newsss.docx", "output.docx")
