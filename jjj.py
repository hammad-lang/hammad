from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx import Document
from openai import OpenAI
from flask import Flask

app=Flask(__name__)

path = 'animation_newsss.docx'
doc = Document(path)
new_doc=Document()

client = OpenAI(
    api_key="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",
)

def chat_with_openai(prompt):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}]
    )
    chatbot_response = response.choices[0].message.content
    return chatbot_response.strip()

def extract_paragraph_style(paragraph):
    style = {
        'font_name': paragraph.style.font.name,
        'font_size': paragraph.style.font.size,
        'spacing_before': paragraph.paragraph_format.space_before,
        'left_indentation': paragraph.paragraph_format.left_indent,
    }
    return style


def apply_paragraph_style(paragraph, style):
    paragraph.style.font.name = style['font_name']
    if style['font_size'] is not None:
        paragraph.style.font.size = style['font_size']
    if style['spacing_before'] is not None:
        paragraph.paragraph_format.space_before =style['spacing_before']
    if style['left_indentation'] is not None:
        paragraph.paragraph_format.left_indent = style['left_indentation']


def iter_block_items(parent):
    if isinstance(parent, _Document):
        print("This is parnet_element_body",parent.element.body)
        parent_elm = parent.element.body

    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def simplifyDocx():
    para_index=0
    final=[]

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):

            if block.text=="":
                pass

            elif block.style.name.startswith("Heading") or block.style.name=="Title":
                style_to_apply = extract_paragraph_style(block)
                new_para=new_doc.add_paragraph(block.text,"Heading 1")
                apply_paragraph_style(new_para, style_to_apply)

            else:
                style_to_apply = extract_paragraph_style(block)
                if para_index==0:
                    prompt = (f'''I am giving you text convert it to simpler english: this is the text : {block.text}
                    
                    ''')
                    text = chat_with_openai(prompt)

                else:
                    prompt = (
                        f'''I am giving you text convert it to simpler english 
                        while ensuring coherence with the previous text, 
                        The system should maintain the original meaning of the content 
                        while employing simpler vocabulary and sentence structures. 
                        this is the text : {block.text}
                        this is previous text:{final[para_index-1]}
                        ''')
                    text = chat_with_openai(prompt)
                final.append(text)


                if block.style.name.startswith("List"):
                    new_para=new_doc.add_paragraph(text,"List Bullet")
                    apply_paragraph_style(new_para, style_to_apply)

                else:
                    new_para=new_doc.add_paragraph(text,block.style.name)
                    apply_paragraph_style(new_para, style_to_apply)

        elif isinstance(block, Table):
            table_data=[]

            for row in block.rows:
                rows=[]
                for cell in row.cells:
                    print(cell.text)
                    rows.append(cell.text)
                table_data.append(rows)

            table = new_doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.style = 'TableGrid'
            for row in range(len(table_data)):
                for col in range(len(table_data[0])):
                    cell = table.cell(row, col)
                    cell.text = f"{table_data[row][col]}"

@app.route('/simplifyDocx',methods=['Get','Post'])
def process():
    simplifyDocx()
    new_doc.save("output.docx")
    return {"Simplified Docx":new_doc.paragraphs[2].text}




if __name__=='__main__':
    app.run(debug=True)
